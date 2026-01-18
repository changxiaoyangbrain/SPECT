import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def register_chinese_font():
    fonts_to_try = [
        ("SimSun", "C:\\Windows\\Fonts\\simsun.ttc"),
        ("SimHei", "C:\\Windows\\Fonts\\simhei.ttf"),
        ("MsYaHei", "C:\\Windows\\Fonts\\msyh.ttf")
    ]
    for font_name, font_path in fonts_to_try:
        if os.path.exists(font_path):
            try:
                if font_path.endswith(".ttf"):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
            except:
                continue
    # Fallback
    path = "C:\\Windows\\Fonts\\simhei.ttf"
    if os.path.exists(path):
        pdfmetrics.registerFont(TTFont("SimHei", path))
        return "SimHei"
    return "Helvetica"

class ReportContent:
    def __init__(self):
        self.title = "SPECT 图像重建实验报告"
        self.date = datetime.now().strftime("%Y年%m月%d日")
        self.sections = [
            {
                "title": "1. 系统矩阵建模",
                "subsections": [
                    {
                        "subtitle": "(a) 建模原理与计算过程",
                        "content": "本实验采用基于射线驱动（Ray-driven）的几何投影模型进行系统矩阵建模。系统矩阵 H 描述了从三维图像空间到二维投影空间的线性映射关系。对于每个投影角度，假设探测器接收到的光子沿直线传播，忽略散射和衰减效应。计算过程中，通过遍历图像体素，计算每个体素中心投影到探测器平面的位置，并利用线性插值将权重分配给相邻的探测器单元。这种方法计算效率高，且能较好地近似物理投影过程。",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 视野离散化设置",
                        "content": "为了实现数值计算，将连续的成像视野（Field of View, FOV）离散化为三维网格。具体参数设置如下：\n- 重建矩阵大小：128 × 128 × 128\n- 体素物理尺寸：3.30 mm × 3.30 mm × 3.30 mm\n此离散化方案与探测器的物理参数（128×128 像素阵列，像素大小 3.30 mm）相匹配，确保了重建空间分辨率与采集系统的一致性。",
                        "image": None
                    },
                    {
                        "subtitle": "(c) 准直器响应模型讨论（可选）",
                        "content": "本实验目前的系统矩阵仅考虑了理想几何投影，即假设准直器孔径无限小，点源在探测器上形成理想的点投影。然而，实际物理系统中，平行孔准直器存在距离模糊效应（Distance-dependent blurring），即点扩展函数（PSF）的宽度随源到准直器距离的增加而线性增加。忽略这一效应是目前重建结果中 SSIM 指标（约 0.54）未能达到极高水平的主要原因之一。若在系统矩阵中引入高斯模糊核建模准直器响应，理论上可显著提升图像的空间分辨率和边缘锐度。",
                        "image": None
                    }
                ]
            },
            {
                "title": "2. OSEM 重建",
                "subsections": [
                    {
                        "subtitle": "(a) 算法原理与参数选择",
                        "content": "实验采用有序子集期望最大化（OSEM）算法。OSEM 是 MLEM 算法的加速版本，其核心思想是将投影数据分组为有序子集，在一次完整迭代中多次更新图像估计。迭代公式为：\n(此处描述 OSEM 更新公式逻辑)\n\n参数选择：\n- 子集数目 (Subsets): 4。选择 4 个子集可以在保证收敛稳定性的同时，将计算速度提升约 4 倍。\n- 迭代次数 (Iterations): 10。经过实验观察，10 次迭代后似然函数值趋于稳定，且图像噪声尚未过度放大，是平衡收敛度与噪声的最佳选择。",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 重建结果展示",
                        "content": "利用上述算法对 Proj.dat 数据进行重建，得到的 MyRecon 结果如下图所示。图像清晰地展示了放射性示踪剂在人体内的三维分布，解剖结构可辨。",
                        "image": r"d:\SPECT\pictures\viz_compare_raw_axial.png",
                        "caption": "图 1: OSEM 重建结果 (MyRecon) 与参考结果的轴向切片对比"
                    },
                    {
                        "subtitle": "(c) 算法扩展性分析（可选）",
                        "content": "OSEM 算法虽然收敛速度快，但由于其基于最大似然估计，随着迭代次数增加，图像的高频噪声会逐渐被放大（Checkerboard effect）。为了克服这一缺点，可以引入最大后验概率（MAP）重建算法。MAP 算法在似然函数的基础上增加了先验概率项（如二次平滑先验或全变分先验），作为正则化约束，能够有效地抑制噪声并保留边缘信息。虽然本实验未实现 MAP，但它是提升低计数 SPECT 图像质量的重要方向。",
                        "image": None
                    }
                ]
            },
            {
                "title": "3. 图像分析评估",
                "subsections": [
                    {
                        "subtitle": "(a) 评估指标说明",
                        "content": "本实验选取了两个客观评价指标：\n1. 均方根误差 (RMSE): 衡量重建图像与参考图像之间像素强度的平均偏差，数值越小表示越接近参考值。\n2. 结构相似性 (SSIM): 从亮度、对比度和结构三个维度衡量两幅图像的相似度，取值范围 [0, 1]，数值越大表示结构越相似。",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 结果定量对比",
                        "content": "将原始重建结果 (MyRecon) 与参考标准 (OSEMReconed.dat) 进行定量对比，结果如下：\n- RMSE: 0.209455\n- SSIM: 0.537552\n分析：RMSE 较低说明整体像素值分布准确。SSIM 约为 0.54，差异主要来源于系统矩阵中未包含准直器模糊效应，导致细节恢复与参考结果（可能使用了更精细的模型）存在差异。",
                        "image": None
                    },
                    {
                        "subtitle": "(c) 后处理效果评估",
                        "content": "为了抑制重建噪声，对 MyRecon 结果进行了三维高斯滤波（FWHM=10mm）。滤波后的结果 (MyFiltered) 与参考滤波结果对比：\n- RMSE: 0.128543 (显著降低)\n- SSIM: 0.329922\n虽然滤波降低了 RMSE，但也平滑了部分结构细节导致 SSIM 下降。这说明后处理需要在去噪和细节保留之间寻找平衡。",
                        "image": r"d:\SPECT\pictures\viz_compare_filtered_axial.png",
                        "caption": "图 2: 滤波后结果 (MyFiltered) 对比"
                    }
                ]
            }
        ]

def generate_pdf(content, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=60, leftMargin=60,
                            topMargin=60, bottomMargin=50)
    Story = []
    cn_font = register_chinese_font()
    styles = getSampleStyleSheet()
    
    # Styles
    style_title = ParagraphStyle(name='TitleCN', parent=styles['Title'], fontName=cn_font, fontSize=18, leading=22, spaceAfter=20, alignment=TA_CENTER)
    style_h1 = ParagraphStyle(name='H1CN', parent=styles['Heading1'], fontName=cn_font, fontSize=14, leading=18, spaceBefore=12, spaceAfter=6)
    style_h2 = ParagraphStyle(name='H2CN', parent=styles['Heading2'], fontName=cn_font, fontSize=12, leading=15, spaceBefore=6, spaceAfter=4)
    style_normal = ParagraphStyle(name='NormalCN', parent=styles['Normal'], fontName=cn_font, fontSize=10.5, leading=16, alignment=TA_JUSTIFY, spaceAfter=6)
    style_caption = ParagraphStyle(name='CaptionCN', parent=styles['Normal'], fontName=cn_font, fontSize=9, leading=12, alignment=TA_CENTER, textColor=colors.grey)

    # Title
    Story.append(Paragraph(content.title, style_title))
    Story.append(Paragraph(f"日期: {content.date}", ParagraphStyle(name='Date', parent=style_normal, alignment=TA_CENTER)))
    Story.append(Spacer(1, 20))

    # Sections
    for section in content.sections:
        Story.append(Paragraph(section['title'], style_h1))
        for sub in section['subsections']:
            Story.append(Paragraph(sub['subtitle'], style_h2))
            
            # Text content handling newlines
            lines = sub['content'].split('\n')
            for line in lines:
                if line.strip():
                    Story.append(Paragraph(line, style_normal))
            
            # Image
            if sub['image'] and os.path.exists(sub['image']):
                try:
                    img = Image(sub['image'], width=400, height=130) # Fixed size for simplicity
                    Story.append(Spacer(1, 6))
                    Story.append(img)
                    if 'caption' in sub:
                        Story.append(Paragraph(sub['caption'], style_caption))
                    Story.append(Spacer(1, 12))
                except:
                    pass

    doc.build(Story)
    print(f"PDF Generated: {filename}")

def generate_docx(content, filename):
    doc = Document()
    
    # Style settings could be elaborate, keeping simple
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    
    # Title
    heading = doc.add_heading(content.title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_p = doc.add_paragraph(f"日期: {content.date}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for section in content.sections:
        doc.add_heading(section['title'], level=1)
        
        for sub in section['subsections']:
            doc.add_heading(sub['subtitle'], level=2)
            doc.add_paragraph(sub['content'])
            
            if sub['image'] and os.path.exists(sub['image']):
                try:
                    doc.add_picture(sub['image'], width=Inches(5.5))
                    if 'caption' in sub:
                        caption = doc.add_paragraph(sub['caption'])
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption'
                except:
                    pass
                    
    doc.save(filename)
    print(f"DOCX Generated: {filename}")

if __name__ == "__main__":
    content = ReportContent()
    output_dir = r"d:\SPECT\reports"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    generate_pdf(content, os.path.join(output_dir, "SPECT大作业.pdf"))
    generate_docx(content, os.path.join(output_dir, "SPECT大作业.docx"))
