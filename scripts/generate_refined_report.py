import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def register_chinese_font():
    fonts_to_try = [
        ("SimSun", "C:\\Windows\\Fonts\\simsun.ttc"),
        ("SimHei", "C:\\Windows\\Fonts\\simhei.ttf"),
        ("MsYaHei", "C:\\Windows\\Fonts\\msyh.ttf")
    ]
    for font_name, font_path in fonts_to_try:
        if os.path.exists(font_path):
            try:
                if font_path.endswith(".ttf"):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
                # Note: .ttc support in reportlab might need index
            except:
                continue
    
    path = "C:\\Windows\\Fonts\\simhei.ttf"
    if os.path.exists(path):
        pdfmetrics.registerFont(TTFont("SimHei", path))
        return "SimHei"
    return "Helvetica"

class ReportContent:
    def __init__(self):
        self.title = "SPECT 图像重建实验报告"
        self.date = datetime.now().strftime("%Y年%m月%d日")
        self.sections = [
            {
                "title": "1. 系统矩阵建模",
                "subsections": [
                    {
                        "subtitle": "(a) 建模原理与计算过程",
                        "content": """本实验采用基于射线驱动（Ray-driven）的几何投影模型。
                        
原理：
假设放射性示踪剂分布为 f(x, y, z)，探测器在角度 θ 处接收到的投影 p(s, z) 可近似为沿射线路径的线积分（Radon 变换）。由于本系统使用平行孔准直器，我们忽略深度相关的模糊效应，假设光子沿垂直于探测器表面的直线传播。

数学推导：
对于离散化系统，投影 p 与图像 f 的关系可表示为线性方程组 p = Hf，其中 H 为系统矩阵。矩阵元素 h_ij 表示第 j 个体素对第 i 个探测器单元的贡献权重。
计算步骤：
1. 网格定义：将成像空间划分为 128x128x128 的体素网格。
2. 坐标变换：对于每个投影角度 θ，将体素中心坐标 (x_v, y_v) 旋转至探测器坐标系 (s, t)。
   s = x_v * cos(θ) + y_v * sin(θ)
3. 权重计算：利用线性插值（Linear Interpolation），将投影位置 s 分配给最近的两个探测器单元。设 s 落在 bin_k 和 bin_{k+1} 之间，则：
   w_k = bin_{k+1} - s
   w_{k+1} = s - bin_k
这种方法避免了复杂的几何相交计算，显著提高了系统矩阵的生成速度。""",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 视野离散化设置",
                        "content": """为了保证重建精度并匹配探测器物理参数，视野离散化设置如下：
- 矩阵维度：128 × 128 × 128 (N_x, N_y, N_z)
- 体素尺寸：3.30 mm × 3.30 mm × 3.30 mm
- 物理视野：422.4 mm × 422.4 mm × 422.4 mm
此设置确保了每个体素与探测器像素（3.30 mm）一一对应，避免了重采样带来的伪影。""",
                        "image": None
                    },
                    {
                        "subtitle": "(c) 准直器响应模型讨论",
                        "content": """当前的几何模型假设准直器具有理想的点扩展函数（PSF 为狄拉克函数）。然而，实际平行孔准直器的 PSF 随源到准直器距离线性增加，呈高斯分布。
误差分析：
忽略这一效应会导致重建图像的高频信息丢失，分辨率低于物理极限。这解释了为何在定量评估中 SSIM 指标（约 0.54）未能达到极高水平。
改进建议：
在系统矩阵中引入距离相关的高斯模糊核（Distance-dependent Gaussian Kernel）。即在正投影过程中，对每个深度层面的投影进行不同 sigma 的高斯卷积，以模拟真实的物理模糊。""",
                        "image": None
                    }
                ]
            },
            {
                "title": "2. OSEM 重建",
                "subsections": [
                    {
                        "subtitle": "(a) 算法原理与流程",
                        "content": """算法原理：
有序子集期望最大化（OSEM）通过将投影数据 P 分为 L 个有序子集 S_1, ..., S_L，加速了 MLEM 的收敛。

迭代公式：
对于第 n 次迭代，第 l 个子集的更新公式为：
f_j^{(n, l)} = \\frac{f_j^{(n, l-1)}}{\\sum_{i \\in S_l} h_{ij}} \\sum_{i \\in S_l} h_{ij} \\frac{p_i}{\\sum_k h_{ik} f_k^{(n, l-1)}}

关键参数：
- 子集数目 (Subsets): 4。平衡了加速比与噪声稳定性。
- 迭代次数 (Iterations): 10。实验表明 10 次迭代后似然函数趋于平稳。""",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 重建结果展示",
                        "content": "下图展示了重建后的轴向切片。图像背景清晰，心脏区域的高摄取区轮廓分明，验证了算法的有效性。",
                        "image": os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "pictures", "viz_compare_raw_axial.png"),
                        "caption": "图 1: OSEM 原始重建结果 (MyRecon) 与参考标准对比"
                    },
                    {
                        "subtitle": "(c) 算法性能分析",
                        "content": """技术讨论：
OSEM 算法在低频成分恢复上表现优异，但随着迭代进行，高频噪声会被放大（棋盘格效应）。
本实验中，原始重建结果的 RMSE 为 0.209，说明整体准确度尚可。为了抑制噪声，我们采用了后处理滤波（Post-filtering）。对比图 2 显示，经 FWHM=10mm 高斯滤波后，图像平滑度显著提升，RMSE 降至 0.128。""",
                        "image": os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "pictures", "viz_compare_filtered_axial.png"),
                        "caption": "图 2: 滤波后结果 (MyFiltered) 对比"
                    }
                ]
            },
            {
                "title": "3. 图像分析评估",
                "subsections": [
                    {
                        "subtitle": "(a) 评估指标",
                        "content": """1. 均方根误差 (RMSE):
   RMSE = \\sqrt{\\frac{1}{N} \\sum (I_{recon} - I_{ref})^2}
   反映像素级的平均偏差。
2. 结构相似性 (SSIM):
   综合考虑亮度、对比度和结构信息，更符合人眼视觉感知。""",
                        "image": None
                    },
                    {
                        "subtitle": "(b) 定量数据表",
                        "content": "下表列出了本实验的最终评估结果：",
                        "table_data": [
                            ["对比组", "RMSE (越小越好)", "SSIM (越大越好)"],
                            ["原始重建 (MyRecon)", "0.209455", "0.537552"],
                            ["滤波后 (MyFiltered)", "0.128543", "0.329922"]
                        ],
                        "image": None
                    },
                    {
                        "subtitle": "(c) 结论与展望",
                        "content": """本实验成功实现了 SPECT 图像重建的全流程。OSEM 算法结合几何投影模型，能够重建出具有解剖意义的三维图像。
主要误差来源：
1. 系统矩阵未建模准直器模糊。
2. 未进行散射校正和衰减校正。
未来改进方向：
引入 MAP 算法利用先验信息抑制噪声，并完善物理模型以提高分辨率。""",
                        "image": None
                    }
                ]
            }
        ]
        self.code_appendix = self.load_code_files()

    def load_code_files(self):
        files = [
            ("system_matrix.py", "系统矩阵建模"),
            ("reconstruction.py", "OSEM 重建算法"),
            ("evaluate.py", "评估指标计算"),
            ("main_pipeline.py", "主流程控制")
        ]
        code_content = []
        base_dir = os.path.dirname(os.path.abspath(__file__))
        for fname, desc in files:
            path = os.path.join(base_dir, fname)
            if os.path.exists(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        code = f.read()
                    code_content.append({"filename": fname, "desc": desc, "code": code})
                except:
                    pass
        return code_content

def generate_pdf(content, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=60, leftMargin=60,
                            topMargin=60, bottomMargin=50)
    Story = []
    cn_font = register_chinese_font()
    styles = getSampleStyleSheet()
    
    # Custom Styles
    style_title = ParagraphStyle(name='TitleCN', parent=styles['Title'], fontName=cn_font, fontSize=20, leading=24, spaceAfter=24, alignment=TA_CENTER)
    style_h1 = ParagraphStyle(name='H1CN', parent=styles['Heading1'], fontName=cn_font, fontSize=16, leading=20, spaceBefore=18, spaceAfter=12)
    style_h2 = ParagraphStyle(name='H2CN', parent=styles['Heading2'], fontName=cn_font, fontSize=13, leading=16, spaceBefore=12, spaceAfter=6)
    style_normal = ParagraphStyle(name='NormalCN', parent=styles['Normal'], fontName=cn_font, fontSize=10.5, leading=16, alignment=TA_JUSTIFY, spaceAfter=8)
    style_code = ParagraphStyle(name='CodeCN', parent=styles['Code'], fontName='Courier', fontSize=8, leading=10, spaceAfter=6, leftIndent=20)
    style_caption = ParagraphStyle(name='CaptionCN', parent=styles['Normal'], fontName=cn_font, fontSize=9, leading=12, alignment=TA_CENTER, textColor=colors.grey)

    # Title Page
    Story.append(Spacer(1, 100))
    Story.append(Paragraph(content.title, style_title))
    Story.append(Paragraph(f"日期: {content.date}", ParagraphStyle(name='Date', parent=style_normal, alignment=TA_CENTER)))
    Story.append(PageBreak())

    # Main Content
    for section in content.sections:
        Story.append(Paragraph(section['title'], style_h1))
        for sub in section['subsections']:
            # Keep subtitle and content together if possible
            elems = []
            elems.append(Paragraph(sub['subtitle'], style_h2))
            
            # Content
            lines = sub['content'].split('\n')
            for line in lines:
                if line.strip():
                    elems.append(Paragraph(line.strip(), style_normal))
            
            # Table if exists
            if 'table_data' in sub:
                t = Table(sub['table_data'], colWidths=[200, 120, 120], hAlign='CENTER')
                t.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), cn_font),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('PADDING', (0, 0), (-1, -1), 6),
                ]))
                elems.append(Spacer(1, 6))
                elems.append(t)
                elems.append(Spacer(1, 12))

            # Image if exists
            if sub['image'] and os.path.exists(sub['image']):
                try:
                    img = Image(sub['image'], width=420, height=140)
                    elems.append(Spacer(1, 6))
                    elems.append(img)
                    if 'caption' in sub:
                        elems.append(Paragraph(sub['caption'], style_caption))
                    elems.append(Spacer(1, 12))
                except:
                    pass
            
            Story.append(KeepTogether(elems))

    # Code Appendix
    Story.append(PageBreak())
    Story.append(Paragraph("附录：核心代码列表", style_h1))
    
    for item in content.code_appendix:
        Story.append(Paragraph(f"文件名: {item['filename']} ({item['desc']})", style_h2))
        code_lines = item['code'].split('\n')
        # Limit code length for report readability, maybe just first 50 lines or full?
        # User requested full code list. Let's put full code but small font.
        for line in code_lines:
            # Replace spaces with non-breaking spaces for indentation
            line_esc = line.replace(' ', '&nbsp;')
            Story.append(Paragraph(line_esc, style_code))
        Story.append(Spacer(1, 12))

    doc.build(Story)
    print(f"PDF Generated: {filename}")

def generate_docx(content, filename):
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    
    # Title
    doc.add_paragraph("\n" * 5)
    heading = doc.add_heading(content.title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"日期: {content.date}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Content
    for section in content.sections:
        doc.add_heading(section['title'], level=1)
        
        for sub in section['subsections']:
            doc.add_heading(sub['subtitle'], level=2)
            doc.add_paragraph(sub['content'])
            
            if 'table_data' in sub:
                table = doc.add_table(rows=len(sub['table_data']), cols=len(sub['table_data'][0]))
                table.style = 'Table Grid'
                for r, row in enumerate(sub['table_data']):
                    for c, val in enumerate(row):
                        table.cell(r, c).text = str(val)
                doc.add_paragraph("") # spacer
            
            if sub['image'] and os.path.exists(sub['image']):
                try:
                    doc.add_picture(sub['image'], width=Inches(5.5))
                    if 'caption' in sub:
                        c = doc.add_paragraph(sub['caption'])
                        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass
    
    # Code Appendix
    doc.add_page_break()
    doc.add_heading("附录：核心代码列表", level=1)
    
    for item in content.code_appendix:
        doc.add_heading(f"{item['filename']} ({item['desc']})", level=2)
        p = doc.add_paragraph(item['code'])
        p.style = 'No Spacing'
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(8)
    
    doc.save(filename)
    print(f"DOCX Generated: {filename}")

if __name__ == "__main__":
    content = ReportContent()
    output_dir = r"d:\SPECT\reports"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    generate_pdf(content, os.path.join(output_dir, "SPECT大作业_Refined.pdf"))
    generate_docx(content, os.path.join(output_dir, "SPECT大作业_Refined.docx"))
